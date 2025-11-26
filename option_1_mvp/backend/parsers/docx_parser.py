"""DOCX Parser - Einfache Text-Extraktion"""
from pathlib import Path
from typing import Dict, Any
from docx import Document
from backend.parsers.base_parser import BaseParser
from backend.utils.logger import setup_logger

logger = setup_logger(__name__)


class DOCXParser(BaseParser):
    def parse(self, file_path: Path) -> Dict[str, Any]:
        try:
            metadata = self.get_metadata(file_path)
            doc = Document(str(file_path))
            
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(text_parts)
            
            metadata.update({"paragraph_count": len(doc.paragraphs)})
            logger.info(f"DOCX geparst: {file_path.name}")
            
            return {"text": full_text, "metadata": metadata, "error": None}
            
        except Exception as e:
            logger.error(f"DOCX Parse Error: {file_path.name} - {e}")
            return {"text": "", "metadata": self.get_metadata(file_path), "error": str(e)}
from pathlib import Path
from typing import Dict, Any
from docx import Document
from backend.parsers.base_parser import BaseParser
from backend.utils.logger import setup_logger

logger = setup_logger(__name__)


class DOCXParser(BaseParser):
    """
    Einfacher DOCX-Parser
    Extrahiert nur Text, keine Tabellen oder Formatierung
    """
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse DOCX und extrahiere Text
        
        Args:
            file_path: Pfad zur DOCX-Datei
        
        Returns:
            Dictionary mit text, metadata, error
        """
        try:
            # Metadaten sammeln
            metadata = self.get_metadata(file_path)
            
            # DOCX öffnen
            doc = Document(str(file_path))
            
            # Text von allen Paragraphen extrahieren
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = "\n".join(text_parts)
            
            # Metadaten erweitern
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            })
            
            logger.info(f"✓ DOCX geparst: {file_path.name} ({len(doc.paragraphs)} Absätze, {len(full_text)} Zeichen)")
            
            return {
                "text": full_text,
                "metadata": metadata,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"✗ Fehler beim Parsen von {file_path.name}: {e}")
            return {
                "text": "",
                "metadata": self.get_metadata(file_path),
                "error": str(e)
            }


if __name__ == "__main__":
    # Test
    print("DOCX Parser Test")
    
    test_dir = Path(__file__).parent.parent.parent / "data" / "input"
    test_dir.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(test_dir.glob("*.docx"))
    
    if docx_files:
        parser = DOCXParser()
        result = parser.parse(docx_files[0])
        
        print(f"\n✓ Test erfolgreich!")
        print(f"  Datei: {result['metadata']['filename']}")
        print(f"  Absätze: {result['metadata'].get('paragraph_count', 'N/A')}")
        print(f"  Zeichen: {len(result['text'])}")
        
        if result['text']:
            print(f"\n  Erste 200 Zeichen:")
            print(f"  {result['text'][:200]}...")
    else:
        print(f"  Keine DOCX-Dateien in {test_dir} gefunden")
