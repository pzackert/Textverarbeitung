from docx import Document as DocxDocument
from pathlib import Path
from datetime import datetime
from typing import List

from .base import BaseParser
from .models import Document
from .exceptions import CorruptedFileError, EmptyDocumentError

class DocxParser(BaseParser):
    supported_formats = ['docx']
    
    def parse(self, file_path: str) -> List[Document]:
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            raise CorruptedFileError(f"Cannot open DOCX: {str(e)}")
        
        try:
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            content = '\n'.join(text_parts)
            
            if not content.strip():
                raise EmptyDocumentError("No text extracted from DOCX")
            
            file_size = path.stat().st_size
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "file_size": file_size,
                "modified_date": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            }
            
            return [Document(
                content=content,
                metadata=metadata,
                source_file=str(file_path),
                file_type="docx"
            )]
        except Exception as e:
            if isinstance(e, EmptyDocumentError):
                raise
            raise CorruptedFileError(f"Error processing DOCX: {str(e)}")
