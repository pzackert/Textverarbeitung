#!/usr/bin/env python3
"""
Sample Data Generator for IFB PROFI Platform.
Generates realistic funding applications (PDF, DOCX) and criteria catalogs.
"""
import os
import sys
import json
import random
from pathlib import Path
from datetime import datetime, timedelta
from dataclasses import dataclass, asdict
from typing import List, Dict

# Add project root to path
sys.path.append(str(Path(__file__).parent.parent))

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm

from docx import Document
from docx.shared import Pt, Inches

# --- Data Models ---

@dataclass
class Company:
    name: str
    address: str
    legal_form: str
    founded: int
    employees: int
    industry: str

@dataclass
class Project:
    title: str
    acronym: str
    description: str
    innovation: str
    start_date: str
    duration_months: int
    total_budget: float
    funding_requested: float
    team_members: List[str]

@dataclass
class Application:
    id: str
    date: str
    company: Company
    project: Project

# --- Sample Data Content ---

APPLICATIONS_DATA = [
    Application(
        id="IFB-PROFI-2024-001",
        date="2024-11-15",
        company=Company(
            name="SmartPort Logistics GmbH",
            address="Am Sandtorkai 1, 20457 Hamburg",
            legal_form="GmbH",
            founded=2020,
            employees=15,
            industry="Logistik / Maritime Wirtschaft"
        ),
        project=Project(
            title="Autonome Drohnensteuerung für Container-Inspektion im Hamburger Hafen",
            acronym="DroneInspect",
            description="Entwicklung eines KI-basierten Steuerungssystems für autonome Drohnen zur visuellen Inspektion von Seecontainern auf Beschädigungen. Das System nutzt Computer Vision zur Echtzeit-Erkennung von Rost, Beulen und Rissen.",
            innovation="Einsatz von Edge-AI direkt auf der Drohne für Latenz-freie Navigation in komplexen Hafenumgebungen. Integration in bestehende Terminal-Management-Systeme.",
            start_date="2025-02-01",
            duration_months=18,
            total_budget=450000.0,
            funding_requested=225000.0,
            team_members=["Dr. Thomas Müller (CTO)", "Sarah Schmidt (AI Lead)", "Jan Jansen (Drone Ops)"]
        )
    ),
    Application(
        id="IFB-PROFI-2024-002",
        date="2024-11-20",
        company=Company(
            name="MedTech Innovations AG",
            address="Luruper Chaussee 149, 22761 Hamburg",
            legal_form="AG",
            founded=2018,
            employees=45,
            industry="Medizintechnik / Life Sciences"
        ),
        project=Project(
            title="Personalisierte 3D-Druck-Implantate für die Kieferchirurgie",
            acronym="JawPrint3D",
            description="Entwicklung eines validierten Workflows für die schnelle Fertigung von patientenspezifischen Kieferimplantaten mittels selektivem Laserschmelzen (SLM) unter Verwendung von biokompatiblem Titan.",
            innovation="Reduktion der Fertigungszeit von 2 Wochen auf 48 Stunden durch automatische Segmentierung von CT-Daten und KI-gestütztes Design.",
            start_date="2025-03-01",
            duration_months=24,
            total_budget=850000.0,
            funding_requested=340000.0,
            team_members=["Prof. Dr. med. Weber (Medical Advisor)", "Lisa Koch (Material Science)", "Markus Wolf (Production)"]
        )
    ),
    Application(
        id="IFB-PROFI-2024-003",
        date="2024-11-25",
        company=Company(
            name="GreenEnergy Solutions GmbH",
            address="Harburger Schloßstraße 6, 21079 Hamburg",
            legal_form="GmbH",
            founded=2022,
            employees=8,
            industry="Erneuerbare Energien"
        ),
        project=Project(
            title="Intelligentes Lastmanagement für Quartiersspeicher in Hamburg",
            acronym="SmartGridHH",
            description="Softwareplattform zur Optimierung von dezentralen Batteriespeichern in Wohnquartieren. Ziel ist die Maximierung des Eigenverbrauchs von PV-Strom und die Entlastung des Verteilnetzes.",
            innovation="Nutzung von Reinforcement Learning zur Prognose von Lastspitzen und dynamischen Preisoptimierung unter Berücksichtigung von Wetterdaten.",
            start_date="2025-01-15",
            duration_months=12,
            total_budget=280000.0,
            funding_requested=140000.0,
            team_members=["Julia Fischer (CEO)", "Kevin Bauer (Software Architect)", "Tim Meyer (Data Scientist)"]
        )
    ),
    Application(
        id="IFB-PROFI-2024-004",
        date="2024-12-01",
        company=Company(
            name="AI-Vision Systems UG",
            address="Große Elbstraße 27, 22767 Hamburg",
            legal_form="UG (haftungsbeschränkt)",
            founded=2023,
            employees=3,
            industry="Künstliche Intelligenz / Software"
        ),
        project=Project(
            title="Automatisierte Qualitätskontrolle in der Lebensmittelproduktion",
            acronym="FoodQualityAI",
            description="Entwicklung eines kostengünstigen Kamerasystems mit integrierter KI zur Erkennung von Fremdkörpern in Abfüllanlagen für Getränke.",
            innovation="Einsatz von Few-Shot Learning, um das System mit sehr wenigen Beispielbildern auf neue Produkte zu trainieren.",
            start_date="2025-04-01",
            duration_months=12,
            total_budget=120000.0,
            funding_requested=60000.0,
            team_members=["Alex Wagner (Founder)", "Sophie Klein (Developer)"]
        )
    ),
    Application(
        id="IFB-PROFI-2024-005",
        date="2024-12-03",
        company=Company(
            name="LogisticOptimizer GmbH & Co. KG",
            address="Billstraße 87, 20539 Hamburg",
            legal_form="GmbH & Co. KG",
            founded=2015,
            employees=120,
            industry="Logistik"
        ),
        project=Project(
            title="CO2-neutrale Letzte Meile durch E-Lastenrad-Hubs",
            acronym="GreenMile",
            description="Pilotprojekt zur Errichtung von Mikro-Hubs in Hamburg-Mitte, von denen aus Pakete per E-Lastenrad zugestellt werden. Entwicklung einer Routing-Software speziell für Lastenräder.",
            innovation="Kombination aus physischer Infrastruktur (Hubs) und digitaler Plattform für dynamisches Tourenmanagement unter Berücksichtigung von Ladezuständen.",
            start_date="2025-05-01",
            duration_months=24,
            total_budget=600000.0,
            funding_requested=240000.0,
            team_members=["Frank Becker (Logistics Manager)", "Laura Schulz (Sustainability Officer)", "Tom Richter (IT)"]
        )
    )
]

# --- Generators ---

def generate_pdf(app: Application, output_path: Path):
    """Generate the main application form as PDF."""
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("IFB Hamburg - PROFI Förderung", styles['Title']))
    story.append(Paragraph(f"Antrag auf Gewährung einer Zuwendung", styles['Heading2']))
    story.append(Spacer(1, 1*cm))

    # Meta Data
    data = [
        ["Antragsnummer:", app.id],
        ["Datum:", app.date],
        ["Programm:", "PROFI Standard"],
    ]
    t = Table(data, colWidths=[4*cm, 10*cm])
    t.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    story.append(t)
    story.append(Spacer(1, 1*cm))

    # 1. Antragsteller
    story.append(Paragraph("1. Antragsteller", styles['Heading2']))
    data = [
        ["Unternehmen:", app.company.name],
        ["Anschrift:", app.company.address],
        ["Rechtsform:", app.company.legal_form],
        ["Gründungsjahr:", str(app.company.founded)],
        ["Mitarbeiter:", str(app.company.employees)],
        ["Branche:", app.company.industry],
    ]
    t = Table(data, colWidths=[4*cm, 10*cm])
    t.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,-1), 0.25, colors.lightgrey),
    ]))
    story.append(t)
    story.append(Spacer(1, 1*cm))

    # 2. Vorhaben
    story.append(Paragraph("2. Das Vorhaben", styles['Heading2']))
    story.append(Paragraph(f"<b>Titel:</b> {app.project.title}", styles['Normal']))
    story.append(Paragraph(f"<b>Akronym:</b> {app.project.acronym}", styles['Normal']))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("<b>Kurzbeschreibung:</b>", styles['Normal']))
    story.append(Paragraph(app.project.description, styles['Normal']))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("<b>Innovationsgehalt:</b>", styles['Normal']))
    story.append(Paragraph(app.project.innovation, styles['Normal']))
    story.append(Spacer(1, 1*cm))

    # 3. Finanzierung
    story.append(Paragraph("3. Finanzierung", styles['Heading2']))
    data = [
        ["Gesamtausgaben:", f"{app.project.total_budget:,.2f} €"],
        ["Beantragte Förderung:", f"{app.project.funding_requested:,.2f} €"],
        ["Eigenmittel:", f"{app.project.total_budget - app.project.funding_requested:,.2f} €"],
        ["Förderquote:", f"{(app.project.funding_requested / app.project.total_budget * 100):.1f} %"],
        ["Laufzeit:", f"{app.project.duration_months} Monate (ab {app.project.start_date})"],
    ]
    t = Table(data, colWidths=[6*cm, 6*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
    ]))
    story.append(t)
    story.append(Spacer(1, 1*cm))

    # 4. Team
    story.append(Paragraph("4. Projektteam", styles['Heading2']))
    for member in app.project.team_members:
        story.append(Paragraph(f"• {member}", styles['Normal']))

    doc.build(story)
    print(f"Generated PDF: {output_path}")

def generate_docx(app: Application, output_path: Path):
    """Generate detailed project description as DOCX."""
    doc = Document()
    
    doc.add_heading(f"Projektbeschreibung: {app.project.acronym}", 0)
    
    doc.add_paragraph(f"Antragsteller: {app.company.name}")
    doc.add_paragraph(f"Datum: {app.date}")
    
    doc.add_heading("1. Ausgangslage und Problemstellung", level=1)
    doc.add_paragraph(
        "Der Markt für die adressierte Lösung wächst stetig. Aktuelle Lösungen sind jedoch oft zu teuer oder technisch veraltet. "
        "Unser Unternehmen hat identifiziert, dass hier ein erhebliches Innovationspotenzial besteht."
    )
    
    doc.add_heading("2. Technische Lösungswege", level=1)
    doc.add_paragraph(app.project.description)
    doc.add_paragraph(
        "Technisch setzen wir auf modernste Standards. Die Architektur ist modular aufgebaut und skalierbar. "
        "Besonderes Augenmerk liegt auf der Datensicherheit und der Einhaltung der DSGVO."
    )
    
    doc.add_heading("3. Innovationsgrad", level=1)
    doc.add_paragraph(app.project.innovation)
    doc.add_paragraph(
        "Im Vergleich zum Stand der Technik (State of the Art) geht unser Ansatz deutlich über inkrementelle Verbesserungen hinaus. "
        "Wir erwarten einen Technologiesprung, der uns einen Wettbewerbsvorteil von ca. 2-3 Jahren verschafft."
    )
    
    doc.add_heading("4. Arbeitsplan", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'AP'
    hdr_cells[1].text = 'Bezeichnung'
    hdr_cells[2].text = 'Dauer (Monate)'
    
    aps = [
        ("AP1", "Anforderungsanalyse & Spezifikation", "2"),
        ("AP2", "Prototyp-Entwicklung", "6"),
        ("AP3", "Integration & Testing", "4"),
        ("AP4", "Validierung & Dokumentation", "2"),
    ]
    
    for ap, name, duration in aps:
        row_cells = table.add_row().cells
        row_cells[0].text = ap
        row_cells[1].text = name
        row_cells[2].text = duration

    doc.add_heading("5. Verwertungsplan", level=1)
    doc.add_paragraph(
        "Nach erfolgreichem Abschluss des Projekts ist die Markteinführung innerhalb von 6 Monaten geplant. "
        "Erste Pilotkunden haben bereits Interesse signalisiert (LOI liegen vor)."
    )

    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

def generate_criteria_catalog(output_path: Path):
    """Generate the criteria catalog JSON."""
    criteria = {
        "version": "1.0",
        "last_updated": datetime.now().strftime("%Y-%m-%d"),
        "categories": [
            {
                "id": "innovation",
                "name": "Innovationsgehalt",
                "weight": 0.4,
                "criteria": [
                    {"id": "tech_level", "question": "Geht die Lösung über den Stand der Technik hinaus?", "max_points": 10},
                    {"id": "usp", "question": "Ist ein technisches Alleinstellungsmerkmal erkennbar?", "max_points": 10}
                ]
            },
            {
                "id": "market",
                "name": "Marktpotenzial",
                "weight": 0.3,
                "criteria": [
                    {"id": "market_size", "question": "Ist der adressierbare Markt groß genug?", "max_points": 10},
                    {"id": "competition", "question": "Wurde der Wettbewerb realistisch eingeschätzt?", "max_points": 10}
                ]
            },
            {
                "id": "feasibility",
                "name": "Umsetzbarkeit",
                "weight": 0.3,
                "criteria": [
                    {"id": "team", "question": "Verfügt das Team über die notwendigen Kompetenzen?", "max_points": 10},
                    {"id": "budget", "question": "Ist der Finanzplan plausibel?", "max_points": 10}
                ]
            }
        ]
    }
    
    with open(output_path, 'w') as f:
        json.dump(criteria, f, indent=2, ensure_ascii=False)
    print(f"Generated Criteria: {output_path}")

def main():
    base_dir = Path("data/samples")
    apps_dir = base_dir / "applications"
    criteria_dir = base_dir / "criteria"
    
    # Create directories
    apps_dir.mkdir(parents=True, exist_ok=True)
    criteria_dir.mkdir(parents=True, exist_ok=True)
    
    print("🚀 Generating Sample Data...")
    
    # Generate Applications
    for app in APPLICATIONS_DATA:
        app_dir = apps_dir / f"{app.id}_{app.company.name.replace(' ', '_')}"
        app_dir.mkdir(exist_ok=True)
        
        # Generate PDF
        generate_pdf(app, app_dir / "antrag.pdf")
        
        # Generate DOCX
        generate_docx(app, app_dir / "projektbeschreibung.docx")
        
        # Save Metadata
        with open(app_dir / "metadata.json", 'w') as f:
            json.dump(asdict(app), f, indent=2, ensure_ascii=False)
            
    # Generate Criteria
    generate_criteria_catalog(criteria_dir / "ifb_profi_criteria.json")
    
    print("\n✅ Sample Data Generation Complete!")
    print(f"   Location: {base_dir.absolute()}")

if __name__ == "__main__":
    main()
